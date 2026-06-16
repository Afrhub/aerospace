# Report Generator — Create PDF/DOCX compliance reports

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io


class ReportGenerator:
    """Generate compliance reports (PDF/DOCX) from workflow + mapping."""

    def generate(self, workflow, format_type='docx'):
        """
        Generate a compliance report.

        Args:
            workflow: Workflow object with steps, mapping, gaps
            format_type: 'pdf' or 'docx'

        Returns:
            bytes (report content) or file path
        """
        if format_type == 'docx':
            return self._generate_docx(workflow)
        elif format_type == 'pdf':
            return self._generate_pdf(workflow)
        else:
            raise ValueError(f"Unsupported format: {format_type}")

    def _generate_docx(self, workflow):
        """Generate DOCX report using python-docx."""
        doc = Document()

        # Title
        title = doc.add_heading('ED-324 Compliance Report', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Program info
        doc.add_heading('Program Information', 1)
        doc.add_paragraph(f"Program: {workflow.program_name}")
        doc.add_paragraph(f"ML System: {workflow.ml_system}")
        doc.add_paragraph(f"Date: {workflow.created_at.strftime('%Y-%m-%d')}")

        # Workflow description
        doc.add_heading('ML Development Workflow', 1)
        for step in workflow.workflow_steps:
            if step.get('name'):  # Only include non-empty steps
                doc.add_heading(step['name'], 2)
                doc.add_paragraph(step.get('description', '(No description provided)'))
                if step.get('tools'):
                    doc.add_paragraph(f"Tools: {', '.join(step['tools'])}")

        # ED-324 Mapping
        doc.add_heading('ED-324 Requirement Mapping', 1)
        mapping = workflow.ed324_mapping
        if mapping:
            doc.add_paragraph(f"Coverage: {mapping.get('coverage_percent', 0)}%")
            for step_name, reqs in mapping.get('mapping', {}).items():
                if step_name:  # Only include non-empty steps
                    doc.add_heading(f"{step_name}", 2)
                    if reqs:
                        doc.add_paragraph(f"Covered requirements: {', '.join(reqs)}")
                    else:
                        doc.add_paragraph("No ED-324 requirements mapped to this step")

        # Gaps
        gaps = workflow.gaps
        if gaps and gaps.get('gap_count', 0) > 0:
            doc.add_heading('Gaps Identified', 1)
            doc.add_paragraph(f"Missing {gaps['gap_count']} ED-324 requirement(s):")
            for gap in gaps.get('gap_details', []):
                doc.add_heading(gap['requirement'], 2)
                doc.add_paragraph(gap['details'])

        # Certification readiness
        doc.add_heading('Certification Readiness', 1)
        doc.add_paragraph(
            "This report documents the ML development workflow mapped to ED-324 requirements. "
            "Review gaps and address them before submitting to regulators."
        )

        # Save to BytesIO and return bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_pdf(self, workflow):
        """Generate PDF report using reportlab."""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib import colors
        from io import BytesIO

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        story = []
        styles = getSampleStyleSheet()

        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=0.3*inch,
            alignment=1
        )
        story.append(Paragraph('ED-324 Compliance Report', title_style))
        story.append(Spacer(1, 0.2*inch))

        # Program info
        story.append(Paragraph('<b>Program Information</b>', styles['Heading2']))
        story.append(Paragraph(f"<b>Program:</b> {workflow.program_name}", styles['Normal']))
        story.append(Paragraph(f"<b>ML System:</b> {workflow.ml_system}", styles['Normal']))
        story.append(Paragraph(f"<b>Date:</b> {workflow.created_at.strftime('%Y-%m-%d')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))

        # Workflow description
        story.append(Paragraph('<b>ML Development Workflow</b>', styles['Heading2']))
        for step in workflow.workflow_steps:
            story.append(Paragraph(f"<b>{step['name']}</b>", styles['Heading3']))
            story.append(Paragraph(step['description'], styles['Normal']))
            if step.get('tools'):
                story.append(Paragraph(f"<i>Tools: {', '.join(step['tools'])}</i>", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

        # ED-324 Mapping
        story.append(PageBreak())
        story.append(Paragraph('<b>ED-324 Requirement Mapping</b>', styles['Heading2']))
        mapping = workflow.ed324_mapping
        if mapping:
            coverage = mapping.get('coverage_percent', 0)
            story.append(Paragraph(f"<b>Coverage: {coverage}%</b>", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            for step_name, reqs in mapping.get('mapping', {}).items():
                story.append(Paragraph(f"<b>{step_name}</b>", styles['Heading3']))
                if reqs:
                    story.append(Paragraph(f"Mapped to: {', '.join(reqs)}", styles['Normal']))
                else:
                    story.append(Paragraph("No ED-324 requirements mapped", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))

        # Gaps
        gaps = workflow.gaps
        if gaps and gaps.get('gap_count', 0) > 0:
            story.append(Paragraph('<b>Gaps Identified</b>', styles['Heading2']))
            story.append(Paragraph(f"Missing {gaps['gap_count']} ED-324 requirement(s):", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            for gap in gaps.get('gap_details', []):
                story.append(Paragraph(f"<b>{gap['requirement']}</b>", styles['Heading3']))
                story.append(Paragraph(gap['details'], styles['Normal']))
                story.append(Spacer(1, 0.05*inch))

        # Certification readiness
        story.append(PageBreak())
        story.append(Paragraph('<b>Certification Readiness</b>', styles['Heading2']))
        story.append(Paragraph(
            "This report documents your ML development workflow mapped to ED-324 requirements. "
            "Address identified gaps before submitting to regulators.",
            styles['Normal']
        ))

        doc.build(story)
        return buffer.getvalue()
